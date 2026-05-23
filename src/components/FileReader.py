import os
import json
import uuid
import platform
import subprocess
from datetime import datetime

try:
    import fitz          # PyMuPDF  — PDF extraction
    _FITZ = True
except ImportError:
    _FITZ = False

try:
    from docx import Document as DocxDocument  
    _DOCX = True
except ImportError:
    _DOCX = False

_HERE        = os.path.dirname(os.path.abspath(__file__))
_PROJECT     = os.path.abspath(os.path.join(_HERE, "..", ".."))

SAMPLES_DIR  = os.path.join(_PROJECT, "evidence", "samples")
UPLOADS_DIR  = os.path.join(_PROJECT, "evidence", "uploads")
INDEX_FILE   = os.path.join(UPLOADS_DIR, "index.json")

SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".doc"}
TYPE_MAP = {
    ".txt":  "txt",
    ".md":   "md",
    ".pdf":  "pdf",
    ".docx": "word",
    ".doc":  "word",
}

def _ensure_uploads_dir() -> None:
    os.makedirs(UPLOADS_DIR, exist_ok=True)


def _new_id() -> str:
    return uuid.uuid4().hex[:8]


def _load_index() -> list[dict]:
    if not os.path.exists(INDEX_FILE):
        return []
    try:
        with open(INDEX_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except (json.JSONDecodeError, OSError):
        return []


def _save_index(index: list[dict]) -> None:
    _ensure_uploads_dir()
    with open(INDEX_FILE, "w", encoding="utf-8") as f:
        json.dump(index, f, indent=2, ensure_ascii=False)


def _extract_text(file_path: str, ext: str) -> str:
    ext = ext.lower()
    try:
        if ext in (".txt", ".md"):
            for enc in ("utf-8", "latin-1", "cp1252"):
                try:
                    with open(file_path, "r", encoding=enc) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                return f.read()

        if ext == ".pdf":
            if not _FITZ:
                return "[PDF extraction unavailable — run: pip install pymupdf]"
            doc = fitz.open(file_path)
            pages = [page.get_text() for page in doc]
            doc.close()
            return "\n".join(pages)

        if ext == ".docx":
            if not _DOCX:
                return "[Word extraction unavailable — run: pip install python-docx]"
            doc = DocxDocument(file_path)
            lines = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            lines.append(cell.text.strip())
            return "\n".join(lines)

        if ext == ".doc":
            # Option 1: antiword (best — install separately)
            try:
                import subprocess
                result = subprocess.run(
                    ["antiword", file_path],
                    capture_output=True, text=True, timeout=30
                )
                if result.returncode == 0 and result.stdout.strip():
                    return result.stdout
            except (FileNotFoundError, subprocess.TimeoutExpired):
                pass

            if _DOCX:
                try:
                    doc = DocxDocument(file_path)
                    lines = [p.text for p in doc.paragraphs if p.text.strip()]
                    if lines:
                        return "\n".join(lines)
                except Exception:
                    pass

            import re as _re
            with open(file_path, "rb") as f:
                raw = f.read()
            strings = _re.findall(rb'[ -~]{4,}', raw)
            text = "\n".join(s.decode("ascii", errors="ignore") for s in strings)
            if text.strip():
                return "[NOTE: .doc raw fallback — install antiword for full support]\n\n" + text
            return "[Could not extract .doc — install antiword for .doc support]"

    except Exception as e:
        return f"[Extraction error: {e}]"

    return ""


def _extract_text_from_bytes(data: bytes, ext: str) -> str:
    import tempfile
    ext = ext.lower()
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
            tmp.write(data)
            tmp_path = tmp.name
        return _extract_text(tmp_path, ext)
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.unlink(tmp_path)
            except OSError:
                pass


def _save_extracted_text(file_id: str, text: str) -> str:
    _ensure_uploads_dir()
    out_path = os.path.join(UPLOADS_DIR, f"{file_id}.txt")
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(text)
    return os.path.relpath(out_path, _PROJECT)


def _record_exists(index: list[dict], original_name: str, size_kb: float) -> bool:
    for entry in index:
        if entry["original_name"] == original_name and entry["size_kb"] == size_kb:
            txt_path = os.path.join(UPLOADS_DIR, f"{entry['id']}.txt")
            return os.path.exists(txt_path)
    return False


def _remove_stale_entry(index: list, original_name: str, size_kb: float) -> None:
    to_remove = []
    for i, entry in enumerate(index):
        if entry["original_name"] == original_name and entry["size_kb"] == size_kb:
            txt_path = os.path.join(UPLOADS_DIR, f"{entry['id']}.txt")
            if not os.path.exists(txt_path):
                to_remove.append(i)
    for i in reversed(to_remove):
        index.pop(i)


def _register_file(
    index: list[dict],
    file_id: str,
    original_name: str,
    file_type: str,
    original_path: str,
    saved_path: str,
    size_kb: float,
    source: str,
) -> dict:
    entry = {
        "id":            file_id,
        "original_name": original_name,
        "file_type":     file_type,
        "source":        source,
        "original_path": original_path,
        "saved_path":    saved_path,
        "size_kb":       size_kb,
        "ingested_at":   datetime.now().isoformat(timespec="seconds"),
    }
    index.append(entry)
    return entry


def _walk_directory_for_supported(root_path: str) -> list[str]:
    found = []
    for dirpath, _, filenames in os.walk(root_path):
        for fname in filenames:
            ext = os.path.splitext(fname)[1].lower()
            if ext in SUPPORTED_EXTENSIONS:
                found.append(os.path.join(dirpath, fname))
    return found


def get_sample_files() -> list[dict]:
    if not os.path.exists(SAMPLES_DIR):
        return []

    files = []
    for fname in sorted(os.listdir(SAMPLES_DIR)):
        ext = os.path.splitext(fname)[1].lower()
        if ext in SUPPORTED_EXTENSIONS:
            full_path = os.path.join(SAMPLES_DIR, fname)
            files.append({
                "name":      fname,
                "path":      full_path,
                "size_kb":   round(os.path.getsize(full_path) / 1024, 1),
                "file_type": TYPE_MAP.get(ext, "txt"),
            })
    return files


def read_sample_file(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    return _extract_text(path, ext)


def ingest_sample_files(sample_file_dicts: list) -> list:
    _ensure_uploads_dir()
    index   = _load_index()
    results = []
    for f in sample_file_dicts:
        fname   = f["name"]
        path    = f["path"]
        ext     = os.path.splitext(fname)[1].lower()
        size_kb = f.get("size_kb", round(os.path.getsize(path) / 1024, 1))
        if _record_exists(index, fname, size_kb):
            continue
        _remove_stale_entry(index, fname, size_kb)
        text       = _extract_text(path, ext)
        file_id    = _new_id()
        saved_path = _save_extracted_text(file_id, text)
        entry      = _register_file(index, file_id, fname, TYPE_MAP.get(ext, "txt"), path, saved_path, size_kb, "sample")
        results.append(entry)
    _save_index(index)
    return results


def ingest_uploaded_files(uploaded_files: list) -> list[dict]:
    _ensure_uploads_dir()
    index   = _load_index()
    results = []

    for uf in uploaded_files:
        fname    = uf.name
        ext      = os.path.splitext(fname)[1].lower()

        if ext not in SUPPORTED_EXTENSIONS:
            continue

        data     = uf.read()
        size_kb  = round(len(data) / 1024, 1)

        if _record_exists(index, fname, size_kb):
            continue
        _remove_stale_entry(index, fname, size_kb)

        text       = _extract_text_from_bytes(data, ext)
        file_id    = _new_id()
        saved_path = _save_extracted_text(file_id, text)
        original_path = getattr(uf, "name", fname)
        entry      = _register_file(
            index, file_id, fname, TYPE_MAP.get(ext, "txt"),
            original_path, saved_path, size_kb, source="upload",
        )
        results.append(entry)

    _save_index(index)
    return results

def detectMounts() -> list[dict]:
    system  = platform.system()
    mounts  = []

    if system == "Darwin":  # macOS
        base = "/Volumes"
        if os.path.exists(base):
            for name in os.listdir(base):
                path = os.path.join(base, name)
                if name not in ("Macintosh HD",) and os.path.isdir(path):
                    mounts.append({"label": name, "path": path})

    elif system == "Linux":
        for base in ("/media", "/mnt"):
            if not os.path.exists(base):
                continue
            for top in os.listdir(base):
                top_path = os.path.join(base, top)
                if os.path.isdir(top_path):
                    sub = os.listdir(top_path)
                    if sub:
                        for s in sub:
                            s_path = os.path.join(top_path, s)
                            if os.path.isdir(s_path):
                                mounts.append({"label": s, "path": s_path})
                    else:
                        mounts.append({"label": top, "path": top_path})

    elif system == "Windows":
        for letter in "ABCDEFGHIJKLMNOPQRSTUVWXYZ":
            path = f"{letter}:\\"
            if os.path.exists(path) and letter not in ("C",):
                try:
                    label_out = subprocess.check_output(
                        ["cmd", "/c", f"vol {letter}:"],
                        stderr=subprocess.DEVNULL,
                        text=True,
                    )
                    label = label_out.split("\n")[0].strip()
                except Exception:
                    label = f"Drive {letter}:"
                mounts.append({"label": label, "path": path})

    return mounts


def get_usb_files(usb_path: str) -> list[dict]:
    _ensure_uploads_dir()
    index   = _load_index()
    results = []

    all_paths = _walk_directory_for_supported(usb_path)

    for file_path in all_paths:
        fname   = os.path.basename(file_path)
        ext     = os.path.splitext(fname)[1].lower()
        size_kb = round(os.path.getsize(file_path) / 1024, 1)

        if _record_exists(index, fname, size_kb):
            continue
        _remove_stale_entry(index, fname, size_kb)

        text       = _extract_text(file_path, ext)
        file_id    = _new_id()
        saved_path = _save_extracted_text(file_id, text)
        entry      = _register_file(
            index, file_id, fname, TYPE_MAP.get(ext, "txt"),
            file_path, saved_path, size_kb, source="usb",
        )
        results.append(entry)

    _save_index(index)
    return results


def load_tracking_index() -> list[dict]:
    return _load_index()


def get_upload_text(file_id: str) -> str:
    index = _load_index()
    for entry in index:
        if entry["id"] == file_id:
            full_path = os.path.join(_PROJECT, entry["saved_path"])
            if os.path.exists(full_path):
                with open(full_path, "r", encoding="utf-8", errors="replace") as f:
                    return f.read()
            break
    txt_path = os.path.join(UPLOADS_DIR, f"{file_id}.txt")
    if not os.path.exists(txt_path):
        return ""
    with open(txt_path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()

def EmptyUploads() -> int:
    if not os.path.exists(UPLOADS_DIR):
        return 0
    count = 0
    for fname in os.listdir(UPLOADS_DIR):
        if fname == ".gitkeep":
            continue
        fpath = os.path.join(UPLOADS_DIR, fname)
        if os.path.isfile(fpath):
            os.unlink(fpath)
            count += 1
    return count